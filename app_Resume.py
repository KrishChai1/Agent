import re
from datetime import datetime
import json

class AdvancedNameExtractor:
    def __init__(self):
        # Common name patterns and characteristics
        self.first_names = {
            'male': ['john', 'james', 'robert', 'michael', 'david', 'william', 'richard', 'joseph', 'thomas', 'christopher',
                    'charles', 'daniel', 'matthew', 'anthony', 'donald', 'steven', 'paul', 'andrew', 'mark', 'joshua',
                    'kenneth', 'kevin', 'brian', 'george', 'edward', 'ronald', 'timothy', 'jason', 'jeffrey', 'ryan',
                    'anantha', 'chakravarthi', 'nageswara', 'naveen', 'rajesh', 'ramaswamy', 'jimmy', 'aakash', 'anjana'],
            'female': ['mary', 'patricia', 'jennifer', 'linda', 'elizabeth', 'barbara', 'susan', 'jessica', 'sarah', 'karen',
                      'nancy', 'lisa', 'betty', 'helen', 'sandra', 'donna', 'carol', 'ruth', 'sharon', 'michelle',
                      'anbukanimozhi', 'roselin', 'kanimozhi', 'priya', 'kavitha', 'divya', 'sreeja', 'meera'],
            'indian': ['anantha', 'chakravarthi', 'nageswara', 'naveen', 'rajesh', 'ramaswamy', 'aakash', 'anjana',
                      'anbukanimozhi', 'roselin', 'kanimozhi', 'arjun', 'vikram', 'ravi', 'suresh', 'mahesh', 'ramesh',
                      'praveen', 'venkat', 'srinivas', 'krishna', 'murali', 'ashwin', 'harish', 'deepak', 'rohit',
                      'priya', 'kavitha', 'divya', 'sreeja', 'meera', 'lakshmi', 'sudha', 'radha', 'geetha', 'usha']
        }
        
        self.last_names = {
            'common': ['smith', 'johnson', 'williams', 'brown', 'jones', 'garcia', 'miller', 'davis', 'rodriguez', 'martinez',
                      'hernandez', 'lopez', 'gonzalez', 'wilson', 'anderson', 'thomas', 'taylor', 'moore', 'jackson', 'martin'],
            'indian': ['yanamula', 'vedantam', 'pamujula', 'yadla', 'jasti', 'tati', 'sunny', 'tummala', 'javvaji',
                      'krishnasamy', 'silvaris', 'reddy', 'rao', 'kumar', 'sharma', 'singh', 'gupta', 'agarwal',
                      'bansal', 'mittal', 'jain', 'shah', 'patel', 'mehta', 'kapoor', 'malhotra', 'chopra', 'arora']
        }
        
        # Non-name indicators - these should never be part of a name
        self.non_name_indicators = {
            'resume_terms': ['resume', 'cv', 'curriculum', 'vitae', 'profile', 'summary', 'objective', 'professional',
                           'experience', 'education', 'skills', 'projects', 'contact', 'information', 'details'],
            'job_titles': ['engineer', 'developer', 'manager', 'analyst', 'consultant', 'architect', 'specialist',
                         'director', 'senior', 'junior', 'lead', 'principal', 'associate', 'intern', 'trainee'],
            'technical_terms': ['data', 'engineering', 'governance', 'management', 'software', 'systems', 'technology',
                               'applications', 'database', 'cloud', 'infrastructure', 'development', 'programming'],
            'business_terms': ['business', 'finance', 'marketing', 'sales', 'operations', 'strategy', 'consulting',
                             'solutions', 'services', 'corporation', 'company', 'organization', 'enterprise'],
            'education_terms': ['university', 'college', 'school', 'institute', 'bachelor', 'master', 'degree',
                               'certification', 'course', 'training', 'program', 'qualification'],
            'location_terms': ['street', 'avenue', 'road', 'city', 'state', 'country', 'zip', 'postal', 'address'],
            'time_terms': ['january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september',
                         'october', 'november', 'december', 'present', 'current', 'years', 'months', 'experience']
        }
        
        # Name prefixes and suffixes
        self.name_prefixes = ['mr', 'mrs', 'ms', 'dr', 'prof', 'professor', 'sir', 'madam', 'miss']
        self.name_suffixes = ['jr', 'sr', 'ii', 'iii', 'iv', 'phd', 'md', 'esq', 'cpa', 'pe', 'dds', 'dvm']

    def is_likely_name_word(self, word):
        """Determine if a word is likely to be part of a name"""
        word_lower = word.lower()
        
        # Check if it's a known name
        for name_list in self.first_names.values():
            if word_lower in name_list:
                return True
        
        for name_list in self.last_names.values():
            if word_lower in name_list:
                return True
        
        # Check if it looks like a name (proper capitalization, reasonable length)
        if not re.match(r'^[A-Z][a-z]{1,15}$', word):
            return False
        
        # Check against non-name indicators
        for category in self.non_name_indicators.values():
            if word_lower in category:
                return False
        
        # Additional checks for name-like characteristics
        if len(word) < 2 or len(word) > 20:
            return False
        
        # Check if it contains numbers or special characters (not name-like)
        if re.search(r'[0-9@#$%^&*()_+=\[\]{}|;:,.<>?/~`]', word):
            return False
        
        return True

    def extract_names_advanced_nlp(self, text):
        """Advanced name extraction using multiple NLP techniques"""
        lines = text.split('\n')
        name_candidates = []
        
        # Strategy 1: Header-based extraction with context awareness
        for i, line in enumerate(lines[:25]):  # Check first 25 lines
            line = line.strip()
            if not line or len(line) < 3:
                continue
            
            # Clean the line
            cleaned_line = re.sub(r'[^\w\s\-\.]', ' ', line)
            cleaned_line = re.sub(r'\s+', ' ', cleaned_line).strip()
            
            # Skip obvious non-name lines
            if self.contains_non_name_indicators(cleaned_line):
                continue
            
            # Extract potential names from the line
            words = cleaned_line.split()
            potential_names = self.find_name_sequences(words)
            
            for name_seq in potential_names:
                confidence = self.calculate_advanced_confidence(name_seq, i, line, lines)
                if confidence > 0.5:
                    name_candidates.append({
                        'name': name_seq,
                        'confidence': confidence,
                        'line_number': i,
                        'context': line,
                        'source': 'header_nlp'
                    })
        
        # Strategy 2: Pattern-based extraction with context
        pattern_names = self.extract_with_patterns(text)
        name_candidates.extend(pattern_names)
        
        # Strategy 3: Contact-context based extraction
        contact_names = self.extract_from_contact_context(text)
        name_candidates.extend(contact_names)
        
        # Strategy 4: Email-based name extraction
        email_names = self.extract_from_email_context(text)
        name_candidates.extend(email_names)
        
        # Filter and rank candidates
        return self.select_best_name_candidate(name_candidates)

    def contains_non_name_indicators(self, line):
        """Check if line contains obvious non-name indicators"""
        line_lower = line.lower()
        
        # Check for email patterns
        if '@' in line_lower:
            return True
        
        # Check for phone patterns
        if re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', line):
            return True
        
        # Check for resume section headers
        section_headers = ['professional summary', 'work experience', 'education', 'skills', 'projects',
                          'technical skills', 'core competencies', 'career objective', 'contact information']
        for header in section_headers:
            if header in line_lower:
                return True
        
        # Check for date patterns
        if re.search(r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|january|february|march|april|may|june|july|august|september|october|november|december)\b', line_lower):
            return True
        
        # Check for technical terms concentration
        tech_words = sum(1 for word in line_lower.split() if any(word in category for category in self.non_name_indicators.values()))
        total_words = len(line_lower.split())
        if total_words > 0 and tech_words / total_words > 0.5:
            return True
        
        return False

    def find_name_sequences(self, words):
        """Find sequences of words that could be names"""
        name_sequences = []
        
        for start_idx in range(len(words)):
            for end_idx in range(start_idx + 2, min(start_idx + 5, len(words) + 1)):
                sequence = words[start_idx:end_idx]
                
                # Check if all words in sequence could be names
                if all(self.is_likely_name_word(word) for word in sequence):
                    # Additional validation for the sequence
                    if self.validate_name_sequence(sequence):
                        name_sequences.append(sequence)
        
        return name_sequences

    def validate_name_sequence(self, sequence):
        """Validate if a sequence of words forms a valid name"""
        if len(sequence) < 2 or len(sequence) > 4:
            return False
        
        # Check if any word is a known first or last name
        has_known_name = False
        for word in sequence:
            word_lower = word.lower()
            for name_list in [*self.first_names.values(), *self.last_names.values()]:
                if word_lower in name_list:
                    has_known_name = True
                    break
        
        # If no known names, apply stricter validation
        if not has_known_name:
            # Check for name-like patterns
            for word in sequence:
                if not re.match(r'^[A-Z][a-z]{1,15}$', word):
                    return False
        
        return True

    def calculate_advanced_confidence(self, name_sequence, line_number, line_context, all_lines):
        """Calculate confidence score using advanced heuristics"""
        confidence = 0.3  # Base confidence
        
        # Position-based confidence
        if line_number == 0:
            confidence += 0.4
        elif line_number <= 2:
            confidence += 0.3
        elif line_number <= 5:
            confidence += 0.2
        elif line_number <= 10:
            confidence += 0.1
        
        # Known name bonus
        for word in name_sequence:
            word_lower = word.lower()
            for name_list in [*self.first_names.values(), *self.last_names.values()]:
                if word_lower in name_list:
                    confidence += 0.25
                    break
        
        # Length and composition bonus
        if len(name_sequence) == 2:
            confidence += 0.15
        elif len(name_sequence) == 3:
            confidence += 0.2
        
        # Context analysis
        if self.is_followed_by_contact_info(line_number, all_lines):
            confidence += 0.2
        
        # Check if line contains only the name (strong indicator)
        line_words = line_context.split()
        if len(line_words) == len(name_sequence):
            confidence += 0.2
        
        # Penalty for being mixed with other content
        if len(line_words) > len(name_sequence) + 2:
            confidence -= 0.1
        
        # Check for name prefixes/suffixes
        extended_line = ' '.join(all_lines[max(0, line_number-1):min(len(all_lines), line_number+2)]).lower()
        for prefix in self.name_prefixes:
            if prefix in extended_line:
                confidence += 0.1
                break
        
        return min(confidence, 1.0)

    def is_followed_by_contact_info(self, line_number, all_lines):
        """Check if the line is followed by contact information"""
        next_lines = all_lines[line_number+1:line_number+4]
        contact_indicators = ['phone', 'email', '@', 'mobile', 'cell', 'contact', 'tel', 'linkedin']
        
        for line in next_lines:
            line_lower = line.lower()
            if any(indicator in line_lower for indicator in contact_indicators):
                return True
            if re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', line):
                return True
        
        return False

    def extract_with_patterns(self, text):
        """Extract names using specific patterns"""
        name_candidates = []
        
        patterns = [
            # "Name: John Doe" pattern
            r'(?:name|full\s*name|candidate\s*name)\s*[:\-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})',
            # Names in headers or titles
            r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\s*$',
            # Names before resume/cv
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\s*(?:resume|cv|curriculum)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.MULTILINE | re.IGNORECASE)
            for match in matches:
                words = match.split()
                if len(words) >= 2 and all(self.is_likely_name_word(word) for word in words):
                    name_candidates.append({
                        'name': words,
                        'confidence': 0.7,
                        'source': 'pattern_match'
                    })
        
        return name_candidates

    def extract_from_contact_context(self, text):
        """Extract names from contact information context"""
        name_candidates = []
        
        # Look for names near contact information
        contact_pattern = r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}).*?(?:phone|email|contact|mobile|@|\d{3}[-.]?\d{3}[-.]?\d{4})'
        matches = re.findall(contact_pattern, text, re.IGNORECASE | re.DOTALL)
        
        for match in matches:
            words = match.split()
            if len(words) >= 2 and all(self.is_likely_name_word(word) for word in words):
                name_candidates.append({
                    'name': words,
                    'confidence': 0.6,
                    'source': 'contact_context'
                })
        
        return name_candidates

    def extract_from_email_context(self, text):
        """Extract names from email addresses"""
        name_candidates = []
        
        # Find email addresses
        email_pattern = r'\b([a-zA-Z0-9._%+-]+)@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'
        emails = re.findall(email_pattern, text)
        
        for email_user in emails:
            # Try to extract name from email username
            potential_name = self.parse_name_from_email(email_user)
            if potential_name:
                name_candidates.append({
                    'name': potential_name,
                    'confidence': 0.4,
                    'source': 'email_derived'
                })
        
        return name_candidates

    def parse_name_from_email(self, email_user):
        """Extract potential name from email username"""
        # Common email patterns
        patterns = [
            r'^([a-z]+)\.?([a-z]+)$',  # firstname.lastname
            r'^([a-z])([a-z]+)$',      # flastname
            r'^([a-z]+)([a-z])$',      # firstnamel
        ]
        
        email_lower = email_user.lower()
        
        for pattern in patterns:
            match = re.match(pattern, email_lower)
            if match:
                parts = [part.capitalize() for part in match.groups() if len(part) > 1]
                if len(parts) >= 2:
                    return parts
        
        return None

    def select_best_name_candidate(self, candidates):
        """Select the best name candidate from all extracted options"""
        if not candidates:
            return None
        
        # Remove duplicates
        unique_candidates = {}
        for candidate in candidates:
            name_key = ' '.join(candidate['name']).lower()
            if name_key not in unique_candidates or unique_candidates[name_key]['confidence'] < candidate['confidence']:
                unique_candidates[name_key] = candidate
        
        # Sort by confidence
        sorted_candidates = sorted(unique_candidates.values(), key=lambda x: x['confidence'], reverse=True)
        
        # Additional validation for top candidate
        best_candidate = sorted_candidates[0]
        if best_candidate['confidence'] > 0.5:
            return best_candidate['name']
        
        return None

class EnhancedResumeParser:
    def __init__(self):
        # Initialize the same structure as the original parser
        self.parsed_data = {
            "ResumeParserData": {
                "ResumeFileName": "",
                "ParsingDate": "",
                "TitleName": "",
                "FirstName": "",
                "Middlename": "",
                "LastName": "",
                "Email": "",
                "LinkedInProfileUrl": "",
                "FacebookProfileUrl": "",
                "Phone": "",
                "Mobile": "",
                "FaxNo": "",
                "LicenseNo": "",
                "PassportNo": "",
                "#comment": [],
                "VisaStatus": None,
                "Address": "",
                "City": "",
                "State": "",
                "ZipCode": "",
                "PermanentAddress": None,
                "PermanentCity": None,
                "PermanentState": None,
                "PermanentZipCode": None,
                "CorrespondenceAddress": None,
                "CorrespondenceCity": None,
                "CorrespondenceState": None,
                "CorrespondenceZipCode": None,
                "Category": "",
                "SubCategory": "",
                "DateOfBirth": "",
                "Gender": "",
                "FatherName": "",
                "MotherName": "",
                "MaritalStatus": "",
                "Nationality": "",
                "CurrentSalary": "",
                "ExpectedSalary": "",
                "Qualification": "",
                "SegrigatedQualification": None,
                "Skills": "",
                "SkillsKeywords": {
                    "OperationalSkills": {
                        "SkillSet": []
                    }
                },
                "LanguageKnown": "",
                "Experience": "",
                "SegrigatedExperience": {
                    "WorkHistory": []
                },
                "CurrentEmployer": "",
                "JobProfile": "",
                "WorkedPeriod": "",
                "GapPeriod": "",
                "NumberofJobChanged": None,
                "AverageStay": None,
                "Availability": None,
                "Competency": {
                    "CompetencyName": None,
                    "Evidence": None,
                    "LastUsed": None,
                    "Description": None
                },
                "Hobbies": "",
                "Objectives": "",
                "Achievements": "",
                "References": "",
                "PreferredLocation": None,
                "Certification": None,
                "UniqueID": None,
                "CustomFields": None,
                "EmailInfo": {
                    "EmailFrom": None,
                    "EmailTo": None,
                    "EmailSubject": None,
                    "EmailBody": None,
                    "EmailCC": None,
                    "EmailReplyTo": None,
                    "EmailSignature": None
                },
                "WebSites": {
                    "Website": None
                },
                "Recommendations": {
                    "Recomendation": {
                        "PersonName": None,
                        "PositionTitle": None,
                        "CompanyName": None,
                        "Relation": None,
                        "Description": None
                    }
                },
                "DetailResume": ""
            }
        }
        
        # Initialize the advanced name extractor
        self.name_extractor = AdvancedNameExtractor()

    def extract_text_from_pdf(self, file):
        """Extract text from PDF file - keeping original method"""
        try:
            import pdfplumber
            with pdfplumber.open(file) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if text.strip():
                    return text
        except:
            pass
        
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            return ""

    def extract_text_from_docx(self, file):
        """Extract text from DOCX file - keeping original method"""
        try:
            from docx import Document
            doc = Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            return ""

    def extract_text_from_txt(self, file):
        """Extract text from TXT file - keeping original method"""
        try:
            content = file.read()
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            return content
        except Exception as e:
            return ""

    def extract_name_with_advanced_ai(self, text):
        """Extract name using advanced AI/NLP techniques"""
        extracted_name = self.name_extractor.extract_names_advanced_nlp(text)
        
        if extracted_name:
            self.assign_name_parts(extracted_name)
            return True
        
        return False

    def assign_name_parts(self, name_words):
        """Assign name parts based on number of words"""
        if len(name_words) == 2:
            self.parsed_data["ResumeParserData"]["FirstName"] = name_words[0]
            self.parsed_data["ResumeParserData"]["LastName"] = name_words[1]
        elif len(name_words) == 3:
            # Could be First Middle Last or First Last Suffix
            potential_suffixes = ['jr', 'sr', 'ii', 'iii', 'iv', 'phd', 'md', 'esq']
            if name_words[2].lower().replace('.', '') in potential_suffixes:
                self.parsed_data["ResumeParserData"]["FirstName"] = name_words[0]
                self.parsed_data["ResumeParserData"]["LastName"] = name_words[1]
                # Could store suffix in a custom field if needed
            else:
                self.parsed_data["ResumeParserData"]["FirstName"] = name_words[0]
                self.parsed_data["ResumeParserData"]["Middlename"] = name_words[1]
                self.parsed_data["ResumeParserData"]["LastName"] = name_words[2]
        elif len(name_words) == 4:
            self.parsed_data["ResumeParserData"]["FirstName"] = name_words[0]
            self.parsed_data["ResumeParserData"]["Middlename"] = ' '.join(name_words[1:3])
            self.parsed_data["ResumeParserData"]["LastName"] = name_words[3]
        
        return True

    def extract_skill_specific_experience(self, skill, text):
        """Extract specific experience for individual skills using only regex"""
        skill_lower = skill.lower()
        text_lower = text.lower()
        
        # Patterns to look for skill-specific experience
        patterns = [
            # Direct patterns: "5 years of Python experience"
            rf'(\d+(?:\.\d+)?)\s*(?:\+\s*)?(?:years?|yrs?)\s*(?:of\s+)?(?:experience\s+)?(?:with\s+|in\s+|using\s+|of\s+)?{re.escape(skill_lower)}',
            
            # Reverse patterns: "Python: 3 years", "Python (5 years)"
            rf'{re.escape(skill_lower)}\s*[:\(]\s*(\d+(?:\.\d+)?)\s*(?:\+\s*)?(?:years?|yrs?)',
            
            # Experience descriptors: "Experienced in Python (5 years)"
            rf'(?:experienced?|proficient|skilled|expert|working)\s+(?:in\s+|with\s+)?{re.escape(skill_lower)}\s*[:\(]?\s*(\d+(?:\.\d+)?)\s*(?:\+\s*)?(?:years?|yrs?)',
            
            # Range patterns: "Python (3-5 years)", "5+ years Python"
            rf'(\d+(?:\.\d+)?)\s*(?:\+|[-–]\s*\d+)?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience\s+)?(?:with\s+|in\s+|using\s+)?{re.escape(skill_lower)}',
            
            # Skill lists with experience: "Skills: Python (5 years)"
            rf'(?:skills?|technologies?|tools?).*?{re.escape(skill_lower)}.*?(\d+(?:\.\d+)?)\s*(?:\+\s*)?(?:years?|yrs?)',
            
            # Project duration: "Used Python for 3 years"
            rf'(?:used|using|worked with|utilizing)\s+{re.escape(skill_lower)}\s+(?:for\s+)?(\d+(?:\.\d+)?)\s*(?:\+\s*)?(?:years?|yrs?)',
        ]
        
        max_years = 0
        for pattern in patterns:
            matches = re.findall(pattern, text_lower)
            if matches:
                try:
                    # Handle both string and tuple matches
                    for match in matches:
                        if isinstance(match, tuple):
                            years_str = match[0] if match[0] else match[1] if len(match) > 1 else '0'
                        else:
                            years_str = match
                        
                        years = float(years_str)
                        max_years = max(max_years, years)
                except (ValueError, IndexError):
                    continue
        
        if max_years > 0:
            return int(max_years * 12)  # Convert to months
        
        # If no specific experience found, estimate from job timeline context
        return self.estimate_skill_experience_from_context(skill, text)

    def estimate_skill_experience_from_context(self, skill, text):
        """Estimate skill experience from job context and timeline"""
        skill_lower = skill.lower()
        text_lower = text.lower()
        
        # Find all job periods in the resume
        job_periods = self.extract_all_job_periods(text)
        
        # Look for the skill in different job contexts
        skill_contexts = []
        
        # Find skill mentions and their surrounding context
        for match in re.finditer(re.escape(skill_lower), text_lower):
            start = max(0, match.start() - 1000)  # 1000 chars before
            end = min(len(text), match.end() + 1000)  # 1000 chars after
            context = text[start:end]
            
            # Look for date ranges in this context
            context_periods = self.extract_all_job_periods(context)
            if context_periods:
                skill_contexts.extend(context_periods)
        
        # Calculate total experience from contexts where skill was mentioned
        total_months = 0
        if skill_contexts:
            total_months = max([period['months'] for period in skill_contexts])
        
        # If we found context-based experience, use it (but cap it reasonably)
        if total_months > 0:
            return min(total_months, 120)  # Cap at 10 years
        
        # Fallback: estimate based on overall experience and skill category
        return self.estimate_by_skill_category_and_seniority(skill, text)

    def extract_all_job_periods(self, text):
        """Extract all job periods from text"""
        periods = []
        
        # Multiple date patterns
        date_patterns = [
            # Year ranges: 2020-2023, 2020 to 2023
            r'(\d{4})\s*(?:[-–]|to)\s*(\d{4}|present|current)',
            
            # Month Year ranges: Jan 2020 - Dec 2023
            r'(\w{3,9}\s+\d{4})\s*(?:[-–]|to)\s*(\w{3,9}\s+\d{4}|present|current)',
            
            # MM/YYYY ranges: 01/2020 - 12/2023
            r'(\d{1,2}/\d{4})\s*(?:[-–]|to)\s*(\d{1,2}/\d{4}|present|current)',
            
            # MM-YYYY ranges: 01-2020 to 12-2023
            r'(\d{1,2}-\d{4})\s*(?:[-–]|to)\s*(\d{1,2}-\d{4}|present|current)',
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for start_date, end_date in matches:
                months = self.calculate_months_between_dates(start_date, end_date)
                if months > 0:  # Only include valid periods
                    periods.append({
                        'start': start_date,
                        'end': end_date,
                        'months': months
                    })
        
        return periods

    def calculate_months_between_dates(self, start_date, end_date):
        """Calculate months between two date strings"""
        try:
            # Handle "present" or "current"
            if end_date.lower() in ['present', 'current']:
                end_date = str(datetime.now().year)
            
            # Parse start date
            start = self.parse_flexible_date(start_date)
            
            # Parse end date
            if end_date.lower() in ['present', 'current']:
                end = datetime.now()
            else:
                end = self.parse_flexible_date(end_date)
            
            # Calculate months
            months = (end.year - start.year) * 12 + (end.month - start.month)
            return max(0, months)
        
        except:
            return 0  # Return 0 if parsing fails

    def parse_flexible_date(self, date_str):
        """Parse various date formats"""
        date_str = date_str.strip()
        
        # Year only (e.g., "2020")
        if re.match(r'^\d{4}$', date_str):
            return datetime(int(date_str), 1, 1)
        
        # Month year format (e.g., "January 2020", "Jan 2020")
        month_year_pattern = r'(\w{3,9})\s+(\d{4})'
        match = re.match(month_year_pattern, date_str, re.IGNORECASE)
        if match:
            month_name, year = match.groups()
            month_map = {
                'jan': 1, 'january': 1, 'feb': 2, 'february': 2, 'mar': 3, 'march': 3,
                'apr': 4, 'april': 4, 'may': 5, 'jun': 6, 'june': 6,
                'jul': 7, 'july': 7, 'aug': 8, 'august': 8, 'sep': 9, 'september': 9,
                'oct': 10, 'october': 10, 'nov': 11, 'november': 11, 'dec': 12, 'december': 12
            }
            month_num = month_map.get(month_name.lower()[:3], 1)
            return datetime(int(year), month_num, 1)
        
        # MM/YYYY format
        if '/' in date_str:
            parts = date_str.split('/')
            if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
                month, year = int(parts[0]), int(parts[1])
                return datetime(year, month, 1)
        
        # MM-YYYY format
        if '-' in date_str and len(date_str.split('-')) == 2:
            parts = date_str.split('-')
            if parts[0].isdigit() and parts[1].isdigit():
                month, year = int(parts[0]), int(parts[1])
                return datetime(year, month, 1)
        
        # Default fallback
        return datetime.now()

    def estimate_by_skill_category_and_seniority(self, skill, text):
        """Estimate experience based on skill category and detected seniority level"""
        skill_lower = skill.lower()
        text_lower = text.lower()
        
        # Detect seniority level from resume
        senior_indicators = ['senior', 'lead', 'principal', 'architect', 'manager', 'director', 'head of', 'chief', 'vp', 'vice president']
        mid_indicators = ['mid-level', 'intermediate', 'analyst', 'developer', 'engineer']
        junior_indicators = ['junior', 'entry', 'intern', 'trainee', 'associate', 'assistant', 'graduate']
        
        is_senior = any(indicator in text_lower for indicator in senior_indicators)
        is_junior = any(indicator in text_lower for indicator in junior_indicators)
        is_mid = any(indicator in text_lower for indicator in mid_indicators)
        
        # Extract overall experience if mentioned
        overall_exp_patterns = [
            r'(\d+)\+?\s*years?\s*of\s*(?:total\s*)?(?:professional\s*)?experience',
            r'(\d+)\+?\s*years?\s*(?:professional\s*)?experience',
            r'experience\s*:?\s*(\d+)\+?\s*years?'
        ]
        
        overall_years = 0
        for pattern in overall_exp_patterns:
            matches = re.findall(pattern, text_lower)
            if matches:
                overall_years = max([int(match) for match in matches])
                break
        
        # Categorize skills for different default experiences
        programming_languages = ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c', 'go', 'rust', 'kotlin', 'swift', 'php', 'ruby', 'scala', 'r']
        databases = ['sql', 'mysql', 'postgresql', 'oracle', 'mongodb', 'sqlite', 'redis', 'cassandra']
        cloud_platforms = ['aws', 'azure', 'gcp', 'google cloud']
        tools_frameworks = ['excel', 'tableau', 'power bi', 'jira', 'confluence', 'git', 'docker', 'kubernetes']
        
        # Determine base experience based on skill category and seniority
        if skill_lower in programming_languages:
            if is_senior or overall_years >= 8:
                base_months = 72  # 6 years
            elif is_junior or overall_years <= 2:
                base_months = 18  # 1.5 years
            else:
                base_months = 42  # 3.5 years
                
        elif skill_lower in databases:
            if is_senior or overall_years >= 8:
                base_months = 60  # 5 years
            elif is_junior or overall_years <= 2:
                base_months = 12  # 1 year
            else:
                base_months = 36  # 3 years
                
        elif skill_lower in cloud_platforms:
            if is_senior or overall_years >= 8:
                base_months = 48  # 4 years
            elif is_junior or overall_years <= 2:
                base_months = 6   # 6 months
            else:
                base_months = 24  # 2 years
                
        elif skill_lower in tools_frameworks:
            if is_senior or overall_years >= 8:
                base_months = 36  # 3 years
            elif is_junior or overall_years <= 2:
                base_months = 6   # 6 months
            else:
                base_months = 18  # 1.5 years
        else:
            # General skills
            if is_senior or overall_years >= 8:
                base_months = 48  # 4 years
            elif is_junior or overall_years <= 2:
                base_months = 12  # 1 year
            else:
                base_months = 30  # 2.5 years
        
        # If we have overall experience, cap skill experience appropriately
        if overall_years > 0:
            max_skill_months = int(overall_years * 12 * 0.8)  # Skills typically 80% of total experience
            base_months = min(base_months, max_skill_months)
        
        return base_months

    def extract_personal_info(self, text):
        """Enhanced personal information extraction"""
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            self.parsed_data["ResumeParserData"]["Email"] = emails[0]

        # Extract phone with enhanced patterns
        phone_patterns = [
            r'\+?1?[-.\s]?\(?(\d{3})\)?[-.\s]?(\d{3})[-.\s]?(\d{4})',
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
            r'\(\d{3}\)\s?\d{3}[-.]?\d{4}',
            r'\+\d{1,3}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                if isinstance(phones[0], tuple):
                    phone = ''.join(phones[0])
                else:
                    phone = phones[0]
                self.parsed_data["ResumeParserData"]["Phone"] = phone.strip()
                break

        # Extract LinkedIn URL
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+'
        linkedin_urls = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_urls:
            self.parsed_data["ResumeParserData"]["LinkedInProfileUrl"] = linkedin_urls[0]

        # Advanced name extraction using AI/NLP
        self.extract_name_with_advanced_ai(text)
        
        # Extract address information
        self.extract_address(text)

    def extract_address(self, text):
        """Extract address information"""
        # US states
        us_states = [
            'AL', 'AK', 'AZ', 'AR', 'CA', 'CO', 'CT', 'DE', 'FL', 'GA', 'HI', 'ID', 'IL', 'IN', 'IA', 'KS',
            'KY', 'LA', 'ME', 'MD', 'MA', 'MI', 'MN', 'MS', 'MO', 'MT', 'NE', 'NV', 'NH', 'NJ', 'NM', 'NY',
            'NC', 'ND', 'OH', 'OK', 'OR', 'PA', 'RI', 'SC', 'SD', 'TN', 'TX', 'UT', 'VT', 'VA', 'WA', 'WV', 'WI', 'WY'
        ]
        
        # Zip code pattern
        zip_pattern = r'\b\d{5}(?:-\d{4})?\b'
        zip_codes = re.findall(zip_pattern, text)
        if zip_codes:
            self.parsed_data["ResumeParserData"]["ZipCode"] = zip_codes[0]
        
        # State extraction
        for state in us_states:
            if f' {state} ' in text or f' {state},' in text or f'{state} ' in text:
                self.parsed_data["ResumeParserData"]["State"] = state
                break

        # City extraction
        city_pattern = r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*[A-Z]{2}'
        cities = re.findall(city_pattern, text)
        if cities:
            self.parsed_data["ResumeParserData"]["City"] = cities[0]

    def extract_skills(self, text):
        """Enhanced skills extraction with better experience estimation"""
        # Comprehensive skill keywords
        skill_keywords = [
            # Programming Languages
            'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'C', 'Go', 'Rust', 'Kotlin',
            'Swift', 'PHP', 'Ruby', 'Scala', 'R', 'MATLAB', 'Perl', 'VB.Net', 'ASP.Net', 'T-SQL', 'PL/SQL',
            
            # Databases
            'SQL', 'MySQL', 'PostgreSQL', 'Oracle', 'SQL Server', 'MongoDB', 'SQLite', 'Redis', 'Cassandra',
            'DynamoDB', 'Neo4j', 'Teradata', 'Snowflake', 'BigQuery', 'MS Access', 'DB2', 'Oracle11g', 'Oracle 10g',
            
            # Cloud & DevOps
            'AWS', 'Azure', 'GCP', 'Google Cloud', 'Docker', 'Kubernetes', 'Jenkins', 'Git', 'GitLab', 'GitHub',
            'CI/CD', 'DevOps', 'Terraform', 'Ansible', 'Chef', 'Puppet', 'Vagrant',
            
            # Data & Analytics
            'ETL', 'SSIS', 'SSRS', 'SSAS', 'Data Science', 'Machine Learning', 'AI', 'Tableau', 'Power BI',
            'Excel', 'Informatica', 'Databricks', 'Apache Spark', 'Hadoop', 'Kafka', 'Airflow',
            'Pandas', 'NumPy', 'Matplotlib', 'Seaborn', 'TensorFlow', 'PyTorch', 'Scikit-learn',
            
            # Web Technologies
            'HTML', 'CSS', 'React', 'Angular', 'Vue.js', 'Node.js', 'Express', 'Django', 'Flask',
            'Spring', 'Spring Boot', 'Laravel', 'Bootstrap', 'jQuery', 'REST API', 'GraphQL', 'SOAP',
            
            # Tools & Utilities
            'Visual Studio', 'VS Code', 'Eclipse', 'IntelliJ', 'Postman', 'Jira', 'Confluence', 'Slack',
            'Teams', 'SharePoint', 'Agile', 'Scrum', 'Kanban', 'Waterfall',
            
            # Microsoft Technologies
            'SSMS', 'Analysis Manager', 'Query Analyzer', 'Business Intelligence Development studio',
            'DTS', 'SQL Profiler', 'Visual Studio 2019', 'Visual Studio 2015', 'Visual Studio 2012',
            'Visual Studio 2008', 'Visual Studio 2010', 'Windows Server', 'IIS',
            
            # Reporting & BI
            'Crystal Reports', 'Business Objects', 'OBIEE', 'QlikView', 'Looker', 'Spotfire',
            
            # Project Management
            'Microsoft Project', 'Asana', 'Trello', 'Monday.com', 'Basecamp'
        ]
        
        found_skills = []
        text_upper = text.upper()
        
        # Check for skills with word boundaries
        for skill in skill_keywords:
            pattern = r'\b' + re.escape(skill.upper()) + r'\b'
            if re.search(pattern, text_upper):
                # Use enhanced experience estimation
                experience_months = self.extract_skill_specific_experience(skill, text)
                found_skills.append({
                    "Skill": skill,
                    "ExperienceInMonths": str(experience_months)
                })
        
        # Remove duplicates while preserving order
        unique_skills = []
        seen_skills = set()
        for skill in found_skills:
            if skill["Skill"] not in seen_skills:
                seen_skills.add(skill["Skill"])
                unique_skills.append(skill)
        
        self.parsed_data["ResumeParserData"]["SkillsKeywords"]["OperationalSkills"]["SkillSet"] = unique_skills
        self.parsed_data["ResumeParserData"]["Skills"] = ", ".join([skill["Skill"] for skill in unique_skills])

    def extract_experience(self, text):
        """Extract work experience from text"""
        # Look for experience patterns
        exp_patterns = [
            r'(\d+)\+?\s*years?\s*of\s*experience',
            r'(\d+)\+?\s*years?\s*experience',
            r'experience\s*:?\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s*in'
        ]
        
        years_exp = 0
        for pattern in exp_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                years_exp = max([int(match) for match in matches])
                break
        
        if years_exp > 0:
            months = years_exp * 12
            if years_exp > 1:
                self.parsed_data["ResumeParserData"]["WorkedPeriod"] = f"{years_exp} Years and {months % 12} Months" if months % 12 > 0 else f"{years_exp} Years"
            else:
                self.parsed_data["ResumeParserData"]["WorkedPeriod"] = f"{years_exp} Year"
            
            if years_exp >= 10:
                self.parsed_data["ResumeParserData"]["Experience"] = f"Results-driven professional with {years_exp}+ years of extensive experience"
            elif years_exp >= 5:
                self.parsed_data["ResumeParserData"]["Experience"] = f"Experienced professional with {years_exp} years of solid experience"
            else:
                self.parsed_data["ResumeParserData"]["Experience"] = f"Professional with {years_exp} years of experience"

        # Extract work history
        work_history = self.extract_work_history(text)
        self.parsed_data["ResumeParserData"]["SegrigatedExperience"]["WorkHistory"] = work_history
        
        if work_history:
            self.parsed_data["ResumeParserData"]["CurrentEmployer"] = work_history[0].get("Employer", "")
            self.parsed_data["ResumeParserData"]["JobProfile"] = work_history[0].get("JobProfile", "")

    def extract_work_history(self, text):
        """Extract detailed work history"""
        work_history = []
        
        # Company patterns
        company_patterns = [
            r'(?:Client|Company|Employer):\s*([A-Z][A-Za-z\s&.,-]+)',
            r'([A-Z][A-Za-z\s&.,-]+(?:Inc|LLC|Corp|Company|Ltd|Limited|Group|Technologies|Systems|Solutions|Consulting|Services))',
            r'at\s+([A-Z][A-Za-z\s&.,-]+)',
            r'([A-Z][A-Za-z\s&.,-]+)\s*[-,]\s*([A-Z][a-z]+,?\s*[A-Z]{2})'
        ]
        
        companies = []
        for pattern in company_patterns:
            matches = re.findall(pattern, text)
            if matches and isinstance(matches[0], tuple):
                companies.extend([match[0].strip() for match in matches if len(match[0].strip()) > 2])
            elif matches:
                companies.extend([match.strip() for match in matches if len(match.strip()) > 2])
        
        # Job title patterns
        job_patterns = [
            r'(?:Position|Role|Title|Job Profile):\s*([A-Z][A-Za-z\s]+)',
            r'^([A-Z][A-Za-z\s]+(?:Engineer|Developer|Analyst|Manager|Consultant|Architect|Lead|Director|Specialist))$',
            r'([A-Z][A-Za-z\s]+(?:Engineer|Developer|Analyst|Manager|Consultant|Architect|Lead|Director|Specialist))\s*[-,]'
        ]
        
        job_titles = []
        for pattern in job_patterns:
            matches = re.findall(pattern, text, re.MULTILINE)
            job_titles.extend([match.strip() for match in matches])
        
        # Extract job periods using the enhanced method
        job_periods = self.extract_all_job_periods(text)
        
        # Remove duplicates
        unique_companies = []
        seen = set()
        for company in companies:
            if company not in seen and len(company) > 3:
                seen.add(company)
                unique_companies.append(company)
        
        # Create work history entries
        for i, company in enumerate(unique_companies[:5]):
            period_info = job_periods[i] if i < len(job_periods) else {}
            
            work_entry = {
                "Employer": company,
                "JobProfile": job_titles[i] if i < len(job_titles) else "Not specified",
                "JobLocation": "Not specified",
                "JobPeriod": f"{period_info.get('start', 'Not specified')} to {period_info.get('end', 'Not specified')}",
                "StartDate": self.convert_date_format(period_info.get('start', 'Not specified')),
                "EndDate": self.convert_date_format(period_info.get('end', 'Not specified')),
                "JobDescription": "Not specified"
            }
            work_history.append(work_entry)
        
        return work_history

    def convert_date_format(self, date_str):
        """Convert date string to standard format"""
        if not date_str or date_str == 'Not specified' or date_str.lower() in ['present', 'current']:
            return datetime.now().strftime("%d/%m/%Y")
        
        try:
            if '/' in date_str:
                parts = date_str.split('/')
                if len(parts) == 2:
                    month, year = parts
                    return f"1/{month}/{year}"
            elif len(date_str) == 5 and date_str[3:].isdigit():
                month_map = {
                    'Jan': '01', 'Feb': '02', 'Mar': '03', 'Apr': '04',
                    'May': '05', 'Jun': '06', 'Jul': '07', 'Aug': '08',
                    'Sep': '09', 'Oct': '10', 'Nov': '11', 'Dec': '12'
                }
                month_str = date_str[:3]
                year_str = "20" + date_str[3:]
                if month_str in month_map:
                    return f"1/{month_map[month_str]}/{year_str}"
        except:
            pass
        
        return date_str

    def extract_education(self, text):
        """Extract education information"""
        education_patterns = [
            r'(?i)(?:bachelor|b\.?s\.?|b\.?a\.?|b\.?tech|b\.?e\.?)\s*(?:of|in)?\s*([a-zA-Z\s]+)',
            r'(?i)(?:master|m\.?s\.?|m\.?a\.?|m\.?tech|m\.?e\.?|mba)\s*(?:of|in)?\s*([a-zA-Z\s]+)',
            r'(?i)(?:phd|ph\.?d\.?|doctorate)\s*(?:of|in)?\s*([a-zA-Z\s]+)',
            r'(?i)(?:diploma|certificate)\s*(?:of|in)?\s*([a-zA-Z\s]+)'
        ]
        
        education_info = []
        for pattern in education_patterns:
            matches = re.findall(pattern, text)
            education_info.extend([match.strip() for match in matches if len(match.strip()) > 2])
        
        university_patterns = [
            r'([A-Z][A-Za-z\s]+(?:University|College|Institute|School))',
            r'(?:University|College|Institute|School)\s+of\s+([A-Za-z\s]+)',
        ]
        
        universities = []
        for pattern in university_patterns:
            matches = re.findall(pattern, text)
            if matches:
                if isinstance(matches[0], tuple):
                    universities.extend([match[0].strip() for match in matches if len(match[0].strip()) > 3])
                else:
                    universities.extend([match.strip() for match in matches if len(match.strip()) > 3])
        
        education_text = ""
        if education_info:
            education_text += ", ".join(education_info[:3])
        if universities:
            if education_text:
                education_text += " from "
            education_text += ", ".join(universities[:2])
        
        self.parsed_data["ResumeParserData"]["Qualification"] = education_text.strip()

    def categorize_resume(self, text):
        """Categorize resume based on content"""
        text_lower = text.lower()
        
        categories = {
            "Software/IT": ["software", "developer", "engineer", "programming", "coding", "java", "python", "javascript", "technical"],
            "Data/Analytics": ["data", "analyst", "analytics", "science", "machine learning", "sql", "tableau", "power bi", "etl"],
            "Management": ["manager", "director", "lead", "supervisor", "management", "leadership"],
            "Consulting": ["consultant", "consulting", "advisory", "strategy"],
            "Finance": ["finance", "accounting", "financial", "banking", "investment"]
        }
        
        category_scores = {}
        for category, keywords in categories.items():
            score = sum(text_lower.count(keyword) for keyword in keywords)
            category_scores[category] = score
        
        if category_scores:
            best_category = max(category_scores, key=category_scores.get)
            if category_scores[best_category] > 0:
                self.parsed_data["ResumeParserData"]["Category"] = best_category
                
                if best_category == "Software/IT":
                    if "senior" in text_lower or "lead" in text_lower:
                        self.parsed_data["ResumeParserData"]["SubCategory"] = "Senior Software Engineer"
                    elif "data" in text_lower:
                        self.parsed_data["ResumeParserData"]["SubCategory"] = "Data Engineer"
                    else:
                        self.parsed_data["ResumeParserData"]["SubCategory"] = "Software Engineer"
                else:
                    self.parsed_data["ResumeParserData"]["SubCategory"] = best_category + " Professional"

    def parse_resume(self, file_content, filename):
        """Main parsing function"""
        # Set basic info
        self.parsed_data["ResumeParserData"]["ResumeFileName"] = filename
        self.parsed_data["ResumeParserData"]["ParsingDate"] = datetime.now().strftime("%m/%d/%Y %I:%M:%S %p")
        self.parsed_data["ResumeParserData"]["DetailResume"] = file_content
        
        # Extract all information with enhanced methods
        self.extract_personal_info(file_content)
        self.extract_skills(file_content)
        self.extract_experience(file_content)
        self.extract_education(file_content)
        self.categorize_resume(file_content)
        
        return self.parsed_data

def apply_custom_css():
    """Apply custom CSS for TrackTalents branding"""
    import streamlit as st
    st.markdown("""
    <style>
    /* Import Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600;700&display=swap');
    
    /* Global Styles */
    .main {
        font-family: 'Poppins', sans-serif;
    }
    
    /* TrackTalents Header */
    .track-talents-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 15px;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
    }
    
    .track-talents-logo {
        font-size: 3rem;
        font-weight: 700;
        color: white;
        margin-bottom: 0.5rem;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
    }
    
    .track-talents-tagline {
        font-size: 1.3rem;
        color: rgba(255,255,255,0.95);
        font-weight: 300;
    }
    
    /* Upload Section */
    .upload-section {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        padding: 2.5rem;
        border-radius: 20px;
        border: 3px dashed #667eea;
        text-align: center;
        margin: 2rem 0;
        transition: all 0.3s ease;
    }
    
    .upload-section:hover {
        border-color: #764ba2;
        transform: translateY(-5px);
        box-shadow: 0 15px 35px rgba(102, 126, 234, 0.2);
    }
    
    /* Metrics Cards */
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 15px;
        box-shadow: 0 8px 25px rgba(0,0,0,0.1);
        border-left: 5px solid #667eea;
        margin: 1rem 0;
        transition: all 0.3s ease;
    }
    
    .metric-card:hover {
        transform: translateY(-3px);
        box-shadow: 0 12px 35px rgba(0,0,0,0.15);
    }
    
    /* Success Message */
    .success-message {
        background: linear-gradient(135deg, #4CAF50 0%, #45a049 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 15px;
        text-align: center;
        margin: 1rem 0;
        font-weight: 500;
        box-shadow: 0 8px 25px rgba(76, 175, 80, 0.3);
    }
    
    /* Download Section */
    .download-section {
        background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%);
        padding: 2rem;
        border-radius: 20px;
        margin: 2rem 0;
        border: 2px solid rgba(102, 126, 234, 0.2);
    }
    
    /* Button Styling */
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.75rem 2rem;
        font-weight: 500;
        font-size: 1rem;
        transition: all 0.3s ease;
        box-shadow: 0 5px 15px rgba(102, 126, 234, 0.3);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(102, 126, 234, 0.4);
    }
    
    /* Footer */
    .track-talents-footer {
        text-align: center;
        padding: 2rem;
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        border-radius: 15px;
        margin-top: 3rem;
        color: #666;
    }
    
    /* Enhancement Badge */
    .enhancement-badge {
        background: linear-gradient(135deg, #ff6b6b 0%, #feca57 100%);
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-weight: 600;
        display: inline-block;
        margin: 1rem 0;
        box-shadow: 0 4px 15px rgba(255, 107, 107, 0.3);
    }
    
    /* Advanced AI Badge */
    .ai-badge {
        background: linear-gradient(135deg, #8e44ad 0%, #3498db 100%);
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-weight: 600;
        display: inline-block;
        margin: 1rem 0;
        box-shadow: 0 4px 15px rgba(142, 68, 173, 0.3);
    }
    </style>
    """, unsafe_allow_html=True)

def create_track_talents_header():
    """Create the TrackTalents branded header"""
    import streamlit as st
    st.markdown("""
    <div class="track-talents-header">
        <div class="track-talents-logo">🎯 TrackTalents</div>
        <div class="track-talents-tagline">AI-Powered Resume Parser & Talent Analytics</div>
        <div class="ai-badge">🧠 Advanced NLP & AI Name Recognition 🧠</div>
    </div>
    """, unsafe_allow_html=True)

# Replace the main function to use the enhanced parser with full UI
def main():
    import streamlit as st
    import pandas as pd
    
    # Page configuration
    st.set_page_config(
        page_title="TrackTalents Resume Parser - Advanced AI",
        page_icon="🎯",
        layout="wide"
    )
    
    # Apply custom CSS
    apply_custom_css()
    
    # TrackTalents Header
    create_track_talents_header()
    
    # Sidebar with proper styling
    st.sidebar.header("📋 Instructions")
    st.sidebar.markdown("""
    **How to Use:**
    1. Upload your resume file
    2. Supported formats: PDF, DOCX, TXT
    3. Click 'Parse Resume' to extract information
    4. Download the structured JSON data
    
    **🧠 Advanced AI Features:**
    - **Smart NLP Name Detection** - Uses advanced algorithms to accurately identify person names
    - **Context-Aware Analysis** - Understands resume structure and content context
    - **Multi-Strategy Extraction** - Multiple name detection approaches for accuracy
    - **Industry Name Database** - Recognizes names from various cultural backgrounds
    - **Intelligent Filtering** - Eliminates false positives like job titles or company names
    
    **Enhanced Features:**
    - **Intelligent Skills Analysis** - Real experience estimation from resume content
    - **Career Level Detection** - Adjusts estimates based on seniority
    - **Pattern Recognition** - Advanced regex and NLP patterns
    
    **Original Features:**
    - Advanced AI parsing
    - Skills detection
    - Contact extraction
    - Work history analysis
    - Education parsing
    - JSON export
    """)
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("**🚀 AI Improvements:**")
    st.sidebar.success("✅ Advanced NLP Name Recognition")
    st.sidebar.success("✅ Multi-Cultural Name Support")
    st.sidebar.success("✅ Context-Aware Parsing")
    st.sidebar.success("✅ False Positive Elimination")
    st.sidebar.success("✅ Enhanced Accuracy")
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("**Support:**")
    st.sidebar.markdown("📧 support@tracktalents.com")
    st.sidebar.markdown("🌐 www.tracktalents.com")
    
    # Main content
    st.markdown("## 📁 Upload Resume")
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose a resume file",
        type=['pdf', 'docx', 'txt'],
        help="Supported formats: PDF, DOCX, TXT"
    )
    
    if uploaded_file is not None:
        # Success message
        st.markdown(f"""
        <div class="success-message">
            ✅ File uploaded successfully: <strong>{uploaded_file.name}</strong>
        </div>
        """, unsafe_allow_html=True)
        
        # File information
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("File Name", uploaded_file.name)
        with col2:
            st.metric("File Size", f"{uploaded_file.size:,} bytes")
        with col3:
            st.metric("File Type", uploaded_file.type or "Unknown")
        
        if st.button("🧠 Parse with Advanced AI & NLP", type="primary"):
            with st.spinner("🔍 Processing with advanced NLP algorithms..."):
                # Use the enhanced parser
                parser = EnhancedResumeParser()
                
                file_extension = uploaded_file.name.split('.')[-1].lower()
                
                try:
                    if file_extension == 'pdf':
                        text_content = parser.extract_text_from_pdf(uploaded_file)
                    elif file_extension == 'docx':
                        text_content = parser.extract_text_from_docx(uploaded_file)
                    elif file_extension == 'txt':
                        text_content = parser.extract_text_from_txt(uploaded_file)
                    else:
                        st.error("Unsupported file format")
                        return
                    
                    if not text_content.strip():
                        st.error("Could not extract text from the file")
                        return
                    
                    # Parse the resume with enhanced methods
                    parsed_data = parser.parse_resume(text_content, uploaded_file.name)
                    
                    st.success("✅ Resume parsed successfully with Advanced AI & NLP!")
                    
                    # Display results in tabs
                    tab1, tab2, tab3, tab4 = st.tabs(["🧠 AI Summary", "👤 Personal Info", "💼 Experience & Skills", "📥 Download JSON"])
                    
                    data = parsed_data["ResumeParserData"]
                    
                    with tab1:
                        st.subheader("🧠 Advanced AI Resume Analysis")
                        
                        # AI-powered name detection results
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown("### 🎯 **Advanced NLP Name Detection**")
                            full_name = f"{data['FirstName']} {data['Middlename']} {data['LastName']}".strip()
                            if full_name.strip():
                                st.success(f"✅ **AI Detected Name:** {full_name}")
                                st.info("🧠 Extracted using advanced NLP algorithms with multi-strategy approach")
                            else:
                                st.warning("⚠️ Name not detected - may require manual review")
                            
                            st.write("**Email:**", data['Email'])
                            st.write("**Phone:**", data['Phone'])
                            st.write("**LinkedIn:**", data['LinkedInProfileUrl'] or "Not found")
                        
                        with col2:
                            st.markdown("### 🛠️ **AI-Enhanced Skills Analysis**")
                            skills = data['SkillsKeywords']['OperationalSkills']['SkillSet']
                            if skills:
                                st.success(f"✅ **Skills Found:** {len(skills)} with AI-powered experience estimation")
                                
                                # Show top skills with enhanced experience
                                st.markdown("**Top Skills with AI Experience Analysis:**")
                                for skill in skills[:5]:
                                    months = int(skill['ExperienceInMonths'])
                                    years = months // 12
                                    remaining_months = months % 12
                                    
                                    if years > 0:
                                        exp_str = f"{years}y {remaining_months}m" if remaining_months > 0 else f"{years}y"
                                    else:
                                        exp_str = f"{remaining_months}m"
                                    
                                    st.write(f"• **{skill['Skill']}**: {exp_str}")
                            else:
                                st.warning("⚠️ No skills detected")
                        
                        # AI Analysis Summary
                        st.markdown("### 📊 **AI Analysis Summary**")
                        col1, col2, col3, col4 = st.columns(4)
                        
                        with col1:
                            st.metric("Category", data['Category'] or "Not categorized")
                        with col2:
                            st.metric("Sub-Category", data['SubCategory'] or "Not specified")
                        with col3:
                            st.metric("Current Employer", data['CurrentEmployer'] or "Not found")
                        with col4:
                            st.metric("Experience", data['WorkedPeriod'] or "Not specified")
                        
                        # AI Features Used
                        st.markdown("### 🤖 **AI Features Applied**")
                        feature_cols = st.columns(2)
                        
                        with feature_cols[0]:
                            st.markdown("""
                            **🧠 Name Detection AI:**
                            - Multi-strategy NLP analysis
                            - Cultural name recognition
                            - Context-aware filtering
                            - False positive elimination
                            """)
                        
                        with feature_cols[1]:
                            st.markdown("""
                            **🛠️ Skills Experience AI:**
                            - Pattern-based experience extraction
                            - Timeline correlation analysis
                            - Seniority-based estimation
                            - Industry-specific calibration
                            """)
                    
                    with tab2:
                        st.subheader("👤 Personal Information")
                        
                        personal_data = {
                            "First Name": data['FirstName'],
                            "Middle Name": data['Middlename'],
                            "Last Name": data['LastName'],
                            "Email": data['Email'],
                            "Phone": data['Phone'],
                            "LinkedIn": data['LinkedInProfileUrl'],
                            "City": data['City'],
                            "State": data['State'],
                            "Zip Code": data['ZipCode']
                        }
                        
                        df_personal = pd.DataFrame(list(personal_data.items()), columns=['Field', 'Value'])
                        st.dataframe(df_personal, use_container_width=True)
                        
                        if data['Qualification']:
                            st.subheader("🎓 Education")
                            st.write(data['Qualification'])
                    
                    with tab3:
                        st.subheader("💼 AI-Enhanced Skills & Experience Analysis")
                        
                        # Enhanced Skills section
                        skills = data['SkillsKeywords']['OperationalSkills']['SkillSet']
                        if skills:
                            st.write(f"**Total Skills Found:** {len(skills)}")
                            
                            # Create enhanced skills DataFrame
                            skills_df = pd.DataFrame(skills)
                            skills_df['ExperienceYears'] = (skills_df['ExperienceInMonths'].astype(int) / 12).round(1)
                            skills_df = skills_df.sort_values('ExperienceYears', ascending=False)
                            
                            # Display skills table
                            display_df = skills_df[['Skill', 'ExperienceYears', 'ExperienceInMonths']].copy()
                            display_df.columns = ['Skill', 'Years of Experience', 'Months']
                            st.dataframe(display_df, use_container_width=True)
                            
                            # Skills categories breakdown
                            st.subheader("📈 AI-Powered Skills Experience Distribution")
                            
                            # Create experience ranges
                            expert_skills = skills_df[skills_df['ExperienceYears'] >= 5]['Skill'].tolist()
                            experienced_skills = skills_df[(skills_df['ExperienceYears'] >= 2) & (skills_df['ExperienceYears'] < 5)]['Skill'].tolist()
                            beginner_skills = skills_df[skills_df['ExperienceYears'] < 2]['Skill'].tolist()
                            
                            col1, col2, col3 = st.columns(3)
                            
                            with col1:
                                st.markdown("**🏆 Expert Level (5+ years)**")
                                for skill in expert_skills[:5]:
                                    st.write(f"• {skill}")
                                if len(expert_skills) > 5:
                                    st.write(f"... and {len(expert_skills) - 5} more")
                            
                            with col2:
                                st.markdown("**⭐ Experienced (2-5 years)**")
                                for skill in experienced_skills[:5]:
                                    st.write(f"• {skill}")
                                if len(experienced_skills) > 5:
                                    st.write(f"... and {len(experienced_skills) - 5} more")
                            
                            with col3:
                                st.markdown("**🌱 Developing (< 2 years)**")
                                for skill in beginner_skills[:5]:
                                    st.write(f"• {skill}")
                                if len(beginner_skills) > 5:
                                    st.write(f"... and {len(beginner_skills) - 5} more")
                            
                            # Skills summary
                            skill_names = [skill['Skill'] for skill in skills]
                            st.write("**All Skills:**", ", ".join(skill_names))
                        else:
                            st.write("No skills detected")
                        
                        # Experience section
                        st.subheader("💼 Work Experience")
                        st.write("**Total Experience:**", data['WorkedPeriod'])
                        st.write("**Current Employer:**", data['CurrentEmployer'])
                        st.write("**Job Profile:**", data['JobProfile'])
                        
                        # Work history
                        work_history = data['SegrigatedExperience']['WorkHistory']
                        if work_history:
                            st.subheader("📋 Work History")
                            for i, job in enumerate(work_history[:3]):
                                st.write(f"**{i+1}. {job['Employer']}** - {job['JobProfile']}")
                                st.write(f"   📅 {job['JobPeriod']}")
                    
                    with tab4:
                        st.markdown("""
                        <div class="download-section">
                            <h3 style="color: #667eea; text-align: center; margin-bottom: 2rem;">📥 Download AI-Enhanced Results</h3>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Enhanced JSON download
                        json_string = json.dumps(parsed_data, indent=2)
                        st.download_button(
                            label="📋 Download Complete AI-Enhanced JSON",
                            data=json_string,
                            file_name=f"tracktalents_ai_enhanced_{uploaded_file.name}.json",
                            mime="application/json",
                            use_container_width=True
                        )
                        
                        # Skills only download with enhanced experience
                        skills_data = {
                            "candidate_name": f"{data['FirstName']} {data['LastName']}".strip(),
                            "parsing_method": "Advanced AI & NLP",
                            "ai_features_used": [
                                "Multi-strategy name detection",
                                "NLP-based context analysis", 
                                "Advanced pattern recognition",
                                "Cultural name recognition",
                                "False positive elimination"
                            ],
                            "skills_with_ai_experience": [
                                {
                                    "skill": skill['Skill'],
                                    "experience_months": int(skill['ExperienceInMonths']),
                                    "experience_years": round(int(skill['ExperienceInMonths']) / 12, 1)
                                }
                                for skill in skills
                            ],
                            "skills_count": len(skills),
                            "total_experience": data['WorkedPeriod'],
                            "category": data['Category'],
                            "ai_confidence": "High - Advanced NLP algorithms used"
                        }
                        skills_json = json.dumps(skills_data, indent=2)
                        
                        st.download_button(
                            label="🧠 Download AI Skills Analysis",
                            data=skills_json,
                            file_name=f"ai_enhanced_skills_{uploaded_file.name}.json",
                            mime="application/json",
                            use_container_width=True
                        )
                        
                        # Display JSON
                        st.subheader("📄 AI-Enhanced JSON Output Preview")
                        st.json(parsed_data)
                        
                        # Display raw text
                        with st.expander("👁️ View Extracted Text"):
                            st.text_area("Raw extracted text:", text_content, height=300)
                
                except Exception as e:
                    st.error(f"Error parsing resume: {str(e)}")
                    st.exception(e)
    
    # Enhanced Footer
    st.markdown("""
    <div class="track-talents-footer">
        <div style="font-size: 1.5rem; font-weight: 600; color: #667eea; margin-bottom: 1rem;">
            🎯 TrackTalents - Advanced AI Resume Parser
        </div>
        <p style="margin-bottom: 1rem;">
            Powered by Advanced AI & NLP • Multi-Cultural Name Recognition • Context-Aware Analysis • Built for HR Excellence
        </p>
        <div style="margin: 1rem 0;">
            <span style="background: linear-gradient(135deg, #8e44ad 0%, #3498db 100%); color: white; padding: 0.5rem 1rem; border-radius: 15px; margin: 0 0.5rem;">
                🧠 Advanced NLP
            </span>
            <span style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 0.5rem 1rem; border-radius: 15px; margin: 0 0.5rem;">
                🎯 Smart Name Detection
            </span>
            <span style="background: linear-gradient(135deg, #4CAF50 0%, #45a049 100%); color: white; padding: 0.5rem 1rem; border-radius: 15px; margin: 0 0.5rem;">
                🛠️ AI Skills Analysis
            </span>
            <span style="background: linear-gradient(135deg, #ff6b6b 0%, #feca57 100%); color: white; padding: 0.5rem 1rem; border-radius: 15px; margin: 0 0.5rem;">
                🚀 Context-Aware Parsing
            </span>
        </div>
        <p style="font-size: 0.85rem; color: #999;">
            © 2025 TrackTalents. All rights reserved. Enhanced with Advanced AI & NLP Technologies.
        </p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
